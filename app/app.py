import streamlit as st
import zipfile
import json
import time
from io import BytesIO
from docx import Document
from datetime import datetime

# Session state for user options
if "user_options_disabled" not in st.session_state:
    st.session_state.user_options_disabled = False

# Session state for holding output
if "processed" not in st.session_state:
    st.session_state.processed = False
if "output_zip" not in st.session_state:
    st.session_state.output_zip = None
if "log_output" not in st.session_state:
    st.session_state.log_output = ""
if "log_messages" not in st.session_state:
    st.session_state.log_messages = []
if "total_files" not in st.session_state:
    st.session_state.total_files = 0

# flags
reset_button_visible = False
process_button_visible = False

# Function to disable the user options
def disable_user_options():
    st.session_state.user_options_disabled = True

# Function to enable the user options
def enable_user_options():
    st.session_state.user_options_disabled = False

# reset everything
def reset_app():
    enable_user_options()
    st.session_state.processed = False
    st.session_state.output_zip = None
    st.session_state.log_output = ""
    st.session_state.log_messages = []
    st.session_state.total_files = 0

# Function to get current timestamp with milliseconds
def get_timestamp():
    now = datetime.now()
    return now.strftime("%H:%M:%S.%f")  # Format: HH:MM:SS.mmmmmm

# Function to show reset button
def show_reset_button():
    global reset_button_visible
    
    if not reset_button_visible:
        reset_button_visible = True
        func_button_col2.button("⬅️ Restart (Clears Results)", on_click=reset_app)

# Function to show process button
def show_process_button():
    global process_button_visible
    
    if not process_button_visible:
        process_button_visible = True
        with func_button_col1:
            process_button_internal = st.button("🚀 Process File", on_click=disable_user_options, type="primary", disabled=st.session_state.user_options_disabled)
        
        return process_button_internal

# Initialize log message list
log_messages = []

# Function to log messages with timestamp
def log(msg: str):
    timestamp = get_timestamp()
    full_msg = f"[{timestamp}] {msg}"
    st.session_state.log_messages.append(full_msg)

st.set_page_config(page_title="Xoul AI Data Converter", layout="centered")

st.title("📄 Xoul AI Data Converter")

st.subheader("About")
st.markdown("This tool converts your Xoul AI exported data to other formats. This makes it easier to use your data on other platforms.")

st.subheader("Privacy")
st.markdown("Your data stays **100% private**: none of your content is stored or logged. No human or machine can review your content.")

st.header("Get Started!")

st.subheader("Step 1: Choose Options", divider=True)

include_markdown = st.checkbox("Include Markdown (.md) output", value=True, disabled=st.session_state.user_options_disabled)
include_word = st.checkbox("Include Word (.docx) output", value=True, disabled=st.session_state.user_options_disabled)

st.subheader("Step 2: Upload File", divider=True)
uploaded_file = st.file_uploader("Upload your Xoul AI Data Export Zip file", type=["zip"])

# Initialize log message list
log_messages = []

if uploaded_file is not None:
    st.success("✅ ZIP file uploaded!")

    st.subheader("Step 3: Start Processing", divider=True)
    st.markdown("To reset the form, click the Restart button")

    # Show process button
    func_button_col1, func_button_col2 = st.columns(2)
    
    process_button = show_process_button()

    if process_button and not st.session_state.processed:
        
        show_reset_button();

        st.subheader("Step 4: Wait for Processing to Finish", divider=True)        
        # Progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            zip_bytes = BytesIO(uploaded_file.read())

            # Create a new in-memory zip to hold output files
            output_zip_buffer = BytesIO()

            with zipfile.ZipFile(zip_bytes, "r") as zip_file:
                file_list = [name for name in zip_file.namelist() if name.endswith(".json") and not name.endswith("/")]
                st.session_state.total_files = len(file_list)

                if st.session_state.total_files == 0:
                    st.warning("No JSON files found in the ZIP.")
                else:
                    st.subheader("📁 Contents of ZIP:")
                    for name in zip_file.namelist():
                        st.text(name)

                    # Open output ZIP
                    with zipfile.ZipFile(output_zip_buffer, "w", zipfile.ZIP_DEFLATED) as output_zip:
                        for idx, name in enumerate(file_list, start=1):
                            try:
                                with zip_file.open(name) as file:
                                    data = json.load(file)
                                    log(f"Loaded: {name}")

                                    st.markdown(f"### 📄 File: `{name}`")

                                    # Generate Markdown
                                    if include_markdown:
                                        markdown_output = json.dumps(data, indent=2)
                                        md_filename = name.replace(".json", ".md")
                                        output_zip.writestr(md_filename, markdown_output)
                                        log(f"Markdown added to ZIP: {md_filename}")

                                    # Generate Word
                                    if include_word:
                                        doc = Document()
                                        doc.add_heading("Generated Document", 0)
                                        doc.add_paragraph(json.dumps(data, indent=2))
                                        doc_buffer = BytesIO()
                                        doc.save(doc_buffer)
                                        doc_filename = name.replace(".json", ".docx")
                                        output_zip.writestr(doc_filename, doc_buffer.getvalue())
                                        log(f"Word doc added to ZIP: {doc_filename}")

                            except Exception as e:
                                error_msg = f"Error processing {name}: {e}"
                                log(error_msg)

                            # Update progress bar
                            progress = idx / st.session_state.total_files
                            progress_bar.progress(progress)
                            status_text.text(f"Processed {idx} of {st.session_state.total_files} files...")

                # Finish writing to ZIP
                output_zip_buffer.seek(0)

                # Done message
                # TODO: add success and error files count
                status_text.success(f"🎉 {st.session_state.total_files}/{st.session_state.total_files} files processed!")
                log(f"{st.session_state.total_files}/{st.session_state.total_files} files processed.")

                # save results
                st.session_state.output_zip = output_zip_buffer.getvalue()
                st.session_state.processed = True
                st.session_state.log_output = "\n".join(st.session_state.log_messages)

        except Exception as e:
            st.error(f"❌ Failed to open ZIP: {e}")
            log(f"Failed to open ZIP: {e}")

    if st.session_state.processed:            
        show_reset_button();

        st.subheader("Step 4: Wait for Processing to Finish", divider=True)        
        # Progress bar
        progress_bar = st.progress(100)
        status_text = st.success(f"🎉 {st.session_state.total_files}/{st.session_state.total_files} files processed!")

        st.subheader("Step 5: Download Results", divider=True)
        
        # Show final download
        output_download_timestamp = get_timestamp()
        st.download_button("📦 Download All Output Files (ZIP)", st.session_state.output_zip, file_name=f"{output_download_timestamp}_converted_xoul_data.zip")
        log("Final ZIP download ready.")

        # Activity log section
        st.subheader("📜 Debug Log", divider=True)
        st.markdown("Send this to Vector if you run into problems.")

        # Display the download button to save the log
        log_download_timestamp = get_timestamp()
        st.download_button("📥 Download Log", st.session_state.log_output, file_name=f"{log_download_timestamp}_xoul_convert_log.txt") 

        # Show log output with timestamps in a scrollable code block
        st.code(st.session_state.log_output, language="bash")

        st.subheader("Page suddenly scrolled to the bottom? It's a bug. Sorry about that! Scroll up ⬆️ to find your files again.", divider=False)