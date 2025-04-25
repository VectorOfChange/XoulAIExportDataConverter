# platform_specific_generators/word/word_xoulai_persona.py
from docx.document import Document as DocxDocument

from globals.globals import NO_DATA_DESCRIPTION
from models.platform_xoulai.persona_xoulai import PersonaPromptSpecXoulAI, PersonaXoulAI

# Persona Helpers
def word_xoulai_add_persona_to_doc(doc: DocxDocument, persona: PersonaXoulAI):
    doc.add_heading(persona.name or "Unnamed Persona", level=2)

    for field_name, value in vars(persona).items():
        if field_name == "name":
            continue

        formatted_name = field_name.replace('_', ' ').title()

        if isinstance(value, PersonaPromptSpecXoulAI):
            doc.add_heading("Prompt Spec", level=3)
            for sub_field, sub_value in vars(value).items():
                doc.add_heading(sub_field.replace('_', ' ').title(), level=4)
                doc.add_paragraph(str(sub_value) if sub_value else NO_DATA_DESCRIPTION)

        else:
            doc.add_heading(formatted_name, level=3)
            
            # Handle the field content
            if value:
                if isinstance(value, list):
                    doc.add_paragraph(", ".join(str(item) for item in value))
                else:
                    doc.add_paragraph(str(value))
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

def word_xoulai_add_all_personas_to_doc(doc: DocxDocument, personas: list[PersonaXoulAI]):
    doc.add_page_break()
    doc.add_heading("Personas", level=1)
    
    if personas:
        for idx, persona in enumerate(personas):
            word_xoulai_add_persona_to_doc(doc, persona)
            
            if idx < len(personas) - 1:
                doc.add_page_break()
    else:
        doc.add_paragraph("No personas were found in the Xoul Data Export ZIP file.")
