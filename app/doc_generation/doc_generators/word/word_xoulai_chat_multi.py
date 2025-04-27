# platform_specific_generators/word/word_xoulai_chat_multi.py
from io import BytesIO
from os import path
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement 

from models.platform_xoulai.chat_multi_xoulai import ChatMultiMessageXoulAI
from doc_generation.doc_generators.word.word_common import word_add_info_section_to_doc, word_add_known_bugs_section_to_doc, word_add_title_page_to_doc, word_add_toc_to_doc
from dtos.file_buffer import FileBuffer
from enums.type_group import TypeGroup
from globals.globals import NO_DATA_DESCRIPTION
from models.platform_xoulai.all_data_xoulai import AllDataXoulAI
from models.platform_xoulai.chat_common_xoulai import ChatConversationXoulAI
from models.platform_xoulai.chat_single_xoulai import ChatSingleMessageXoulAI

def word_xoulai_add_chat_multi_info_to_doc(doc: DocxDocument, chat_info: ChatConversationXoulAI):
    doc.add_page_break()
    doc.add_heading("Chat Information", level=2)
    
    doc.add_heading("Name", level=3)
    doc.add_paragraph(chat_info.name or "Unnamed Chat Conversation")

    for field_name, value in vars(chat_info).items():
        if field_name == "name":
            continue  # Already used for the heading

        formatted_name = field_name.replace('_', ' ').title()

        if field_name == "xouls":
            doc.add_heading("Characters", level=3)
            if value:
                for xoul in value:
                    doc.add_heading(xoul.name or "Unnamed Character", level=4)
                    for sub_field, sub_value in vars(xoul).items():
                        if sub_field != "name": # skip name field, it is already used for the heading 
                            sub_formatted_name = sub_field.replace('_', ' ').title()
                            doc.add_heading(sub_formatted_name, level=5)
                            doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        elif field_name == "personas":
            doc.add_heading("Personas", level=3)
            if value:
                for persona in value:
                    doc.add_heading(persona.name or "Unnamed Persona", level=4)
                    for sub_field, sub_value in vars(persona).items():
                        if sub_field != "name": # skip name field, it is already used for the heading
                            sub_formatted_name = sub_field.replace('_', ' ').title()
                            doc.add_heading(sub_formatted_name, level=5)
                            doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        elif field_name == "scenario":
            doc.add_heading("Scenario", level=3)
            if value:
                for scenario in value:
                    doc.add_heading(scenario.name or "Unnamed Scenario", level=4)
                    for sub_field, sub_value in vars(value).items():
                        if sub_field != "name": # skip name field, it is already used for the heading
                            sub_formatted_name = sub_field.replace('_', ' ').title()
                            doc.add_heading(sub_formatted_name, level=4)
                            if isinstance(sub_value, list):
                                if sub_value:
                                    doc.add_paragraph(", ".join(str(item) for item in sub_value))
                                else:
                                    doc.add_paragraph(NO_DATA_DESCRIPTION)
                            else:
                                doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        else:
            doc.add_heading(formatted_name, level=3)
            if value is not None:
                if isinstance(value, list):
                    if value:
                        doc.add_paragraph(", ".join(str(item) for item in value))
                    else:
                        doc.add_paragraph(NO_DATA_DESCRIPTION)
                else:
                    doc.add_paragraph(str(value))
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)
      
def word_xoulai_add_chat_multi_transcript_to_doc(doc: Document, messages: list[ChatMultiMessageXoulAI]):
    """
    Adds a formatted group chat transcript to a Word document.

    Args:
        doc: The Word document to add the transcript to.
        messages: A list of ChatMultiMessageXoulAI objects representing the chat messages.
        conversation_name: The name of the conversation/group chat.
    """
    doc.add_page_break()
    doc.add_heading("Group Chat Transcript", level=1)

    if not messages:
        doc.add_paragraph("No chat messages available.")
        return

    # Define desaturated, light shades of color for different speakers.  Visually distinctive but not distracting.
    color_shades = {
        "#E0E0E0": "Light Grey",
        "#B0B0B0": "Medium Dark Grey",
        "#909090": "Very Dark Grey",
        "#E0FFFF": "Light Cyan",
        "#FFF0F5": "Lavender Blush",
        "#FFFACD": "Lemon Chiffon",
        "#FFDAB9": "Peachpuff",
        "#E0F8E0": "Very Pale Green"
    }

    speaker_colors = {}  # Dictionary to store assigned colors

    def get_speaker_color(speaker_name: str) -> str:
        """
        Assigns a unique shade of grey to each speaker.  Uses a dictionary
        to store previously assigned colors.

        Args:
            speaker_name: The name of the speaker.

        Returns:
            A string representing the RGB color in hexadecimal format (e.g., "F0F0F0").
        """
        if speaker_name not in speaker_colors:
            # If this is a new speaker, assign them a color.
            num_speakers = len(speaker_colors)
            color_index = num_speakers % len(color_shades)  # Cycle through shades
            speaker_colors[speaker_name] = color_shades[color_index]
        return speaker_colors[speaker_name]

    def add_styled_paragraph(doc: Document, text: str, fill_color: str, font_color: RGBColor, align: WD_ALIGN_PARAGRAPH):
        """
        Adds a styled paragraph to the document.

        Args:
            doc: The Word document.
            text: The text of the paragraph.
            fill_color: The background fill color in hexadecimal format (e.g., "F0F0F0").
            font_color: The font color as an RGBColor object.
            align: The paragraph alignment (e.g., WD_ALIGN_PARAGRAPH.LEFT).
        """
        para = doc.add_paragraph()
        para.alignment = align
        run = para.add_run(text)
        run.font.color.rgb = font_color

        # Add shading background to paragraph
        pPr = para.paragraph_format.element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)  # Fill background
        pPr.append(shd)

        return para

    for message in messages:
        speaker_name = message.author_name or "Unknown Speaker"
        content = message.content or NO_DATA_DESCRIPTION
        role = (message.author_type or "unknown").lower() # added role

        if role != "user":
            # Determine styling based on speaker (using the helper function)
            bg_fill = get_speaker_color(speaker_name)
            font_rgb = RGBColor(0, 0, 0)  # Black text for all speakers
            alignment = WD_ALIGN_PARAGRAPH.LEFT  # Default alignment
        else :
            # User color
            bg_fill = "000000"
            font_rgb = RGBColor(255,255,255)
            alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add speaker name (bold and italic)
        name_para = add_styled_paragraph(doc, speaker_name, bg_fill, font_rgb, alignment)
        for run in name_para.runs:
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)

        # Add message content
        add_styled_paragraph(doc, content, bg_fill, font_rgb, alignment)

        # Handle alternate regenerations
        if message.metadata and message.metadata.alternative_regenerations:
            for alt_content in message.metadata.alternative_regenerations:
                alt_para = doc.add_paragraph()
                alt_para.alignment = alignment
                run_label = alt_para.add_run("Alternate Version")
                run_label.italic = True
                run_label.font.color.rgb = RGBColor(0, 0, 0)
                run_label.font.highlight_color = 7

                run_alt = alt_para.add_run("\n" + (alt_content or NO_DATA_DESCRIPTION))
                run_alt.font.color.rgb = font_rgb

                pPr_alt = alt_para.paragraph_format.element.get_or_add_pPr()
                shd_alt = OxmlElement('w:shd')
                shd_alt.set(qn('w:val'), 'clear')
                shd_alt.set(qn('w:color'), 'auto')
                shd_alt.set(qn('w:fill'), bg_fill)
                pPr_alt.append(shd_alt)

def get_character_names_for_filename(character_names: list[str]) -> str:
    """
    Generates a string of character names for use in a filename,
    limiting the total length to a maximum of 50 characters and
    distributing the available space as evenly as possible among the names.

    Args:
        character_names: A list of character names (strings).

    Returns:
        A string containing a truncated and combined version of the character names,
        or an empty string if the input list is empty.
    """
    if not character_names:
        return ""

    max_length = 50
    num_names = len(character_names)

    # Calculate the base length and remaining characters
    base_length = max_length // num_names
    remaining_chars = max_length % num_names

    name_segments = []
    for i, name in enumerate(character_names):
        # Distribute the remaining characters
        extra_char = 1 if i < remaining_chars else 0
        segment_length = base_length + extra_char
        name_segments.append(name[:segment_length])  # Truncate the name

    return "_".join(name_segments)  # Join the segments with underscores

# Example Usage
# character_names = ["Sarah", "Levi", "Loki", "VeryLongNameIndeed", "Shorty"]
# filename_string = get_character_names_for_filename(character_names)
# print(filename_string)  # Output: Sarah_Levi_Loki_VeryLo_Short

def word_xoulai_generate_chat_multi_docs(platform_data: AllDataXoulAI, on_progress=None) -> list[FileBuffer]:
    doc_buffers: list[FileBuffer] = []
    
    for idx, chat_multi in enumerate(platform_data.chats_multi):
        doc = Document(path.join('app', 'assets', 'numbered_heading_template.docx'))

        # Content
        word_add_title_page_to_doc(doc, platform_data.platform, TypeGroup.CHAT, chat_multi.get_chat_description())
        word_add_toc_to_doc(doc)
        word_add_info_section_to_doc(doc, platform_data.platform)
        word_add_known_bugs_section_to_doc(doc)

        word_xoulai_add_chat_multi_info_to_doc(doc, chat_multi.conversation)
        word_xoulai_add_chat_multi_transcript_to_doc(doc, chat_multi.messages)

        # Save document to buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)

        # get filename
        character_names = get_character_names_for_filename(chat_multi.get_character_names())
        persona_names = chat_multi.get_persona_names(return_none_if_empty=True)
        persona_name = persona_names[0] if persona_names else None

        filename = f"XoulAI_Group_Chat_Transcript_{idx}_{character_names}{'_with_' + persona_name if persona_name else ''}.docx"
        
        doc_buffers.append(FileBuffer(doc_buffer, filename))

        if on_progress:
            on_progress()
    
    return doc_buffers