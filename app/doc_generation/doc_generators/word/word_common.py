# platform_specific_generators/word/word_common.py
from docx.document import Document as DocxDocument

from models.platform_xoulai.chat_common_xoulai import ChatConversationXoulAI
from enums.bug_type import BugType
from enums.platform import Platform
from enums.type_group import TypeGroup
from globals.globals import APP_VERSION, KNOWN_BUGS, NO_DATA_DESCRIPTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement 
from docx.shared import RGBColor, Pt

# Common/Shared Helpers
def word_xoulai_add_chat_info_to_doc(doc: DocxDocument, chat_info: ChatConversationXoulAI):
    doc.add_page_break()
    doc.add_heading("Chat Information", level=1)
    
    doc.add_heading("Name", level=2)
    doc.add_paragraph(chat_info.name or "Unnamed Chat Conversation")

    for field_name, value in vars(chat_info).items():
        if field_name == "name":
            continue  # Already used for the heading

        formatted_name = field_name.replace('_', ' ').title()

        if field_name == "xouls":
            doc.add_heading("Characters", level=2)
            if value:
                for xoul in value:
                    doc.add_heading(xoul.name or "Unnamed Character", level=3)
                    for sub_field, sub_value in vars(xoul).items():
                        if sub_field != "name": # skip name field, it is already used for the heading
                            sub_formatted_name = sub_field.replace('_', ' ').title()
                            doc.add_heading(sub_formatted_name, level=4)
                            doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        elif field_name == "personas":
            doc.add_heading("Personas", level=2)
            if value:
                for persona in value:
                    doc.add_heading(persona.name or "Unnamed Persona", level=3)
                    for sub_field, sub_value in vars(persona).items():
                        if sub_field != "name": # skip name field, it is already used for the heading
                            sub_formatted_name = sub_field.replace('_', ' ').title()
                            doc.add_heading(sub_formatted_name, level=4)
                            doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        elif field_name == "scenario":
            doc.add_heading("Scenario", level=2)
            if value:
                doc.add_heading(value.name or "Unnamed Scenario", level=3)
                for sub_field, sub_value in vars(value).items():
                    if sub_field != "name": # skip name field, it is already used for the heading
                        sub_formatted_name = sub_field.replace('_', ' ').title()
                        doc.add_heading(sub_formatted_name, level=4)
                        if isinstance(sub_value, list):
                            if sub_value:
                                doc.add_paragraph(", ".join(str(item) for item in sub_value))
                            else:
                                doc.add_paragraph(NO_DATA_DESCRIPTION)
                        else:
                            doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        else:
            doc.add_heading(formatted_name, level=2)
            if value is not None:
                if isinstance(value, list):
                    if value:
                        doc.add_paragraph(", ".join(str(item) for item in value))
                    else:
                        doc.add_paragraph(NO_DATA_DESCRIPTION)
                else:
                    doc.add_paragraph(str(value))
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

def word_add_info_section_to_doc(doc: DocxDocument, platform: Platform):
    # Comes after TOC, so NO page page
    
    doc.add_heading("Introduction", level=1)
    
    # about section
    doc.add_heading("About this Document", level=2)
    doc.add_paragraph("This is the information that is within the Xoul Data Export ZIP file. Icon and picture links have been removed because they were just links to the Xoul AI site, which is now offline.")
    
    # color section
    doc.add_heading("Color in this Document", level=2)
    doc.add_paragraph(f"This document may use colors. This is particularly used in group chat transcripts to highlight the different characters. These colors are very light so they don't strain your eyes while you read.") 
    doc.add_paragraph("If you are using Word in Dark Mode or with a Dark Office Theme, Word will automatically adjust the colors and they will become dark and intense. If you find this difficult to read, try changing your Office Theme to a light theme. You can find this settings under: File --> Account --> Office Theme in modern versions of Word.)")

    # data manipulation section
    doc.add_heading("Data Adjustment/Manipulation", level=2)
    if platform is Platform.XOULAI:
        doc.add_paragraph("The data in this document is in the original Xoul AI format.")
    else:
        doc.add_paragraph(f"The data in this document has been adjusted and/or manipulated to be optimized for use with the {platform} platform.")
    
    doc.add_paragraph("If you generated data for multiple platforms, check the other files in the ZIP file.")
    doc.add_paragraph("If you need data for a different platform, go back to the website and generate it.")

    # slug section
    doc.add_heading("Slugs", level=2)
    doc.add_paragraph("A \"slug\" is a unique ID used to link data together. For example, if a scenario uses a lorebook, the scenario data will include the slug of the lorebook.")
    doc.add_paragraph("If you created the referenced item (for example, you want to go to the lorebook that was included in the scenario, and you created that lorebook), you can search this document for the slug and it should bring you to it.")
    
    # bugs section
    doc.add_heading("Bugs and Technical Issues", level=2)
    p = doc.add_paragraph("Bugs? Errors? Technical issues?")
    p.add_run().add_break()
    p.add_run("Missing Data? Weird computer code in the generated documents?")
    doc.add_paragraph("Report it and I'll fix it! My contact information is in the Contact section. Remember to save the log from the website if possible.")
    
    # contact section
    doc.add_heading("Contact")
    doc.add_paragraph("You can contact me here:")
    doc.add_paragraph("The best way to contact me is to send me a discord DM: @vectorofchange (join the Xoul Discord at https://discord.gg/xoul to get access to DM me)")
    doc.add_paragraph("You can also tag me (@VectorOfChange) in the #alts channel in the Xoul Discord.")
    doc.add_paragraph("If you don't use Discord, you can start an Issue on Github on the source code.")
    p = doc.add_paragraph()
    p.add_run(" Note: ").bold = True
    p.add_run("Anything you post in the Github Repo will be public.")

    # Source code
    doc.add_heading("Source Code")
    doc.add_paragraph("Github Source Code Repo: https://github.com/VectorOfChange/XoulAIExportDataConverter")

    # App Version
    doc.add_heading("App Version")
    doc.add_paragraph(f"App Version: {APP_VERSION}")

def word_add_known_bugs_section_to_doc(doc: DocxDocument):
    doc.add_page_break()
    
    doc.add_heading("Known Bugs and Issues", level=1)
    
    doc.add_paragraph(f"Known bugs and issue with app version {APP_VERSION} that impact data integrity or data quality (bugs that do not impact data are not listed here):")
    
    data_bugs = [bug["description"] for bug in KNOWN_BUGS if bug["type"] == BugType.DATA]

    if data_bugs:
        for bug_description in data_bugs:
            doc.add_paragraph(bug_description, style='List Bullet')
    else:
        doc.add_paragraph("No Known Bugs", style='List Bullet')

def word_add_title_page_to_doc(doc: DocxDocument, platform: Platform, type_group: TypeGroup, description: list[str] = []):
    # Title Page
    prefix = "Converted " if platform is not Platform.XOULAI else ""
    title = doc.add_paragraph(f"{prefix}Xoul AI Data", style='Title')

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(type_group.value, style='Subtitle')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    description_subtitle = doc.add_paragraph(style='Intense Quote')
    description_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if description:
        for idx, description_line in enumerate(description):
            r = description_subtitle.add_run(description_line)
            r.add_break()
            r.add_break()
    
    description_subtitle.add_run("Modified for Platform: ")
    if platform is Platform.XOULAI:
        description_subtitle.add_run("Unmodified (Original Xoul AI format)")
    else:
        description_subtitle.add_run(platform)

    beta_tag = doc.add_paragraph()
    r = beta_tag.add_run("BETA TEST", style='Intense Emphasis')
    r.font.color.rgb = RGBColor(255, 75, 75) # red
    r.font.size = Pt(16)
    beta_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER

def word_add_toc_to_doc(doc: DocxDocument):
    doc.add_page_break()
    # TOC
    doc.add_heading("Table of Contents", 0)
    toc_paragraph = doc.add_paragraph()
    toc_run = toc_paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click on this text, then click Update Field to show the Table of Contents. For accurate page numbers, right click the table of contents again after it is shown, and click Update Field again."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = toc_run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    toc_run.add_break(WD_BREAK.PAGE)