# platform_specific_generators/word/word_xoulai_character.py
from docx.document import Document as DocxDocument

from globals.globals import NO_DATA_DESCRIPTION
from models.platform_xoulai.character_xoulai import CharacterXoulAI

# Character Helpers
def word_xoulai_add_character_to_doc(doc: DocxDocument, character: CharacterXoulAI):
    # Add Heading 1 with the character's name
    doc.add_heading(character.name or "Unnamed Character", level=2)

    # Iterate over the character fields dynamically
    for field_name, value in vars(character).items():
        if field_name == "name":
            continue  # already used in Heading 1

        # Format the field name nicely (e.g., backstory_spec -> Backstory Spec)
        formatted_name = field_name.replace('_', ' ').title()
        doc.add_heading(formatted_name, level=3)

        # Handle the field content
        if value:
            if isinstance(value, list):
                doc.add_paragraph(", ".join(str(item) for item in value))
            else:
                doc.add_paragraph(str(value))
        else:
            doc.add_paragraph(NO_DATA_DESCRIPTION)

def word_xoulai_add_all_characters_to_doc(doc: DocxDocument, characters: list[CharacterXoulAI]):
    doc.add_page_break()
    doc.add_heading("Characters", level=1)
    
    if characters:
        for idx, character in enumerate(characters):
            word_xoulai_add_character_to_doc(doc, character)
            
            if idx < len(characters) - 1:
                doc.add_page_break()
    else:
        doc.add_paragraph("No characters were found in the Xoul Data Export ZIP file.")
