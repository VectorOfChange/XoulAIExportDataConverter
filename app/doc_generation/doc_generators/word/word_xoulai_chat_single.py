# platform_specific_generators/word/word_xoulai_chat_single.py
from io import BytesIO
from os import path
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement 

from doc_generation.doc_generators.word.word_xoulai_common import word_xoulai_add_chat_info_to_doc
from doc_generation.doc_generators.word.word_common import word_add_info_section_to_doc, word_add_known_bugs_section_to_doc, word_add_title_page_to_doc, word_add_toc_to_doc
from dtos.file_buffer import FileBuffer
from enums.type_group import TypeGroup
from globals.globals import NO_DATA_DESCRIPTION
from models.platform_xoulai.all_data_xoulai import AllDataXoulAI
from models.platform_xoulai.chat_common_xoulai import ChatConversationXoulAI
from models.platform_xoulai.chat_single_xoulai import ChatSingleMessageXoulAI

def word_xoulai_add_chat_single_transcript_to_doc(doc: DocxDocument, messages: list[ChatSingleMessageXoulAI], char_name: str):
    doc.add_page_break()
    doc.add_heading("Chat Transcript", level=1)

    if not messages:
        doc.add_paragraph("No chat messages available.")
        return

    for message in messages:
        role = (message.role or "unknown").lower()
        content = message.content or NO_DATA_DESCRIPTION
        name = "User" if role == "user" else char_name #TODO: change the persona name for user if possible

        # Determine styling based on role
        if role == "assistant":
            bg_fill = "F0F0F0"  # Light grey background
            font_rgb = RGBColor(0, 0, 0)  # Black text
            alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif role == "user":
            bg_fill = "000000"  # Black background
            font_rgb = RGBColor(255, 255, 255)  # White text
            alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            bg_fill = "FFFFFF"  # Default white
            font_rgb = RGBColor(0, 0, 0)  # Default black
            alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Helper to add a styled paragraph (with shading and font color)
        def add_styled_paragraph(text: str, fill_color: str, font_color: RGBColor, align):
            para = doc.add_paragraph()
            para.alignment = align
            run = para.add_run(text)
            run.font.color.rgb = font_color

            # Add shading background to paragraph
            ppr = para.paragraph_format.element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill_color)  # Fill background
            ppr.append(shd)

            return para

        # Add speaker name (bold and italic)
        name_para = add_styled_paragraph(name, bg_fill, font_rgb, alignment)
        for run in name_para.runs:
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)  # Set the font size (adjust as needed, default is 10.5)

        # Add message content
        add_styled_paragraph(content, bg_fill, font_rgb, alignment)

        # Handle alternate regenerations if any
        if message.metadata and message.metadata.alternative_regenerations:
            for alt_content in message.metadata.alternative_regenerations:
                # Add "Alternate Version" label
                alt_para = doc.add_paragraph()
                alt_para.alignment = alignment
                run_label = alt_para.add_run("Alternate Version")
                run_label.italic = True
                run_label.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                run_label.font.highlight_color = 7  # Yellow highlight (Word standard value)

                # Add alternate content
                run_alt = alt_para.add_run("\n" + (alt_content or NO_DATA_DESCRIPTION))
                run_alt.font.color.rgb = font_rgb

                # Add shading to entire alternate version paragraph
                ppr_alt = alt_para.paragraph_format.element.get_or_add_pPr()
                shd_alt = OxmlElement('w:shd')
                shd_alt.set(qn('w:val'), 'clear')
                shd_alt.set(qn('w:color'), 'auto')
                shd_alt.set(qn('w:fill'), bg_fill)
                ppr_alt.append(shd_alt)

def word_xoulai_generate_chat_single_docs(platform_data: AllDataXoulAI, on_progress=None) -> list[FileBuffer]:
    doc_buffers: list[FileBuffer] = []
    
    for idx, chat_single in enumerate(platform_data.chats_single):
        doc = Document(path.join('app', 'assets', 'numbered_heading_template.docx'))

        # Content
        word_add_title_page_to_doc(doc, platform_data.platform, TypeGroup.CHAT, chat_single.get_chat_description())
        word_add_toc_to_doc(doc)
        word_add_info_section_to_doc(doc, platform_data.platform)
        word_add_known_bugs_section_to_doc(doc)

        word_xoulai_add_chat_info_to_doc(doc, chat_single.conversation)
        word_xoulai_add_chat_single_transcript_to_doc(doc, chat_single.messages, chat_single.get_character_name())

        # Save document to buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)

        # get filename
        persona_names = chat_single.get_persona_names(return_none_if_empty=True)
        persona_name = persona_names[0] if persona_names else None

        filename = f"XoulAI_Single_Chat_Transcript_{idx}_{chat_single.get_character_slug()}{'_with_' + persona_name if persona_name else ''}.docx"
        
        doc_buffers.append(FileBuffer(doc_buffer, filename))

        if on_progress:
            on_progress()
    
    return doc_buffers