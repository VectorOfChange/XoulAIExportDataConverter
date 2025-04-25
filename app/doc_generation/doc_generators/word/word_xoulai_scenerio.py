# platform_specific_generators/word/word_xoulai_scenario.py
from docx.document import Document as DocxDocument

from globals.globals import NO_DATA_DESCRIPTION
from models.platform_xoulai.scenario_xoulai import ScenarioObjectiveXoulAI, ScenarioPromptSpecXoulAI, ScenarioXoulAI

# Scenario Helpers
def word_xoulai_add_scenario_to_doc(doc: DocxDocument, scenario: ScenarioXoulAI):
    doc.add_heading(scenario.name or "Unnamed Scenario", level=2)

    for field_name, value in vars(scenario).items():
        if field_name == "name":
            continue

        formatted_name = field_name.replace('_', ' ').title()
        if isinstance(value, ScenarioPromptSpecXoulAI):
            doc.add_heading("Prompt Spec", level=3)
            for sub_field, sub_value in vars(value).items():
                doc.add_heading(sub_field.replace('_', ' ').title(), level=4)
                doc.add_paragraph(str(sub_value) if sub_value else NO_DATA_DESCRIPTION)

        elif isinstance(value, ScenarioObjectiveXoulAI):
            doc.add_heading("Objective", level=3)
            doc.add_heading("Description", level=4)
            doc.add_paragraph(value.description or NO_DATA_DESCRIPTION)
            if value.meters:
                doc.add_heading("Meters", level=4)
                for meter in value.meters:
                    doc.add_heading(meter.name or "Unnamed Meter", level=5)
                    doc.add_paragraph(f"Description: {meter.description or 'None/Empty'}")
                    doc.add_paragraph(f"Value: {meter.value}")
            else:
                doc.add_paragraph("No meters")

        else:
            doc.add_heading(formatted_name, level=3)
            if value:
                if isinstance(value, list):
                    doc.add_paragraph(", ".join(str(item) for item in value))
                else:
                    doc.add_paragraph(str(value))
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

def word_xoulai_add_all_scenarios_to_doc(doc: DocxDocument, scenarios: list[ScenarioXoulAI]):
    doc.add_page_break()
    doc.add_heading("Scenarios", level=1)
    
    if scenarios:
        for idx, scenario in enumerate(scenarios):
            word_xoulai_add_scenario_to_doc(doc, scenario)
            
            if idx < len(scenarios) - 1:
                doc.add_page_break()
    else:
        doc.add_paragraph("No scenarios were found in the Xoul Data Export ZIP file.")
