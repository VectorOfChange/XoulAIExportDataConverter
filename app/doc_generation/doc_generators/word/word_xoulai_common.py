# platform_specific_generators/word/word_xoulai_common.py
from docx.document import Document as DocxDocument

from models.platform_xoulai.chat_common_xoulai import ChatConversationXoulAI
from globals.globals import NO_DATA_DESCRIPTION

def word_xoulai_add_chat_info_to_doc(doc: DocxDocument, chat_info: ChatConversationXoulAI):
    doc.add_page_break()
    doc.add_heading("Chat Information", level=1)

    doc.add_heading("Name", level=2)
    doc.add_paragraph(chat_info.name or "Unnamed Chat Conversation")

    for field_name, value in vars(chat_info).items():
        if field_name == "name":
            continue  # Already used for the heading

        formatted_name = field_name.replace('_', ' ').title()

        if field_name == "xouls":
            doc.add_heading("Characters", level=2)
            if value:
                for xoul in value:
                    doc.add_heading(xoul.name or "Unnamed Character", level=3)
                    for sub_field, sub_value in vars(xoul).items():
                        if sub_field != "name": # skip name field, it is already used for the heading
                            sub_formatted_name = sub_field.replace('_', ' ').title()
                            doc.add_heading(sub_formatted_name, level=4)
                            doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        elif field_name == "personas":
            doc.add_heading("Personas", level=2)
            if value:
                for persona in value:
                    doc.add_heading(persona.name or "Unnamed Persona", level=3)
                    for sub_field, sub_value in vars(persona).items():
                        if sub_field != "name": # skip name field, it is already used for the heading
                            sub_formatted_name = sub_field.replace('_', ' ').title()
                            doc.add_heading(sub_formatted_name, level=4)
                            doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        elif field_name == "scenario":
            doc.add_heading("Scenario", level=2)
            if value:
                doc.add_heading(value.name or "Unnamed Scenario", level=3)
                for sub_field, sub_value in vars(value).items():
                    if sub_field != "name": # skip name field, it is already used for the heading
                        sub_formatted_name = sub_field.replace('_', ' ').title()
                        doc.add_heading(sub_formatted_name, level=4)
                        if isinstance(sub_value, list):
                            if sub_value:
                                doc.add_paragraph(", ".join(str(item) for item in sub_value))
                            else:
                                doc.add_paragraph(NO_DATA_DESCRIPTION)
                        else:
                            doc.add_paragraph(str(sub_value) if sub_value is not None else NO_DATA_DESCRIPTION)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)

        else:
            doc.add_heading(formatted_name, level=2)
            if value is not None:
                if isinstance(value, list):
                    if value:
                        doc.add_paragraph(", ".join(str(item) for item in value))
                    else:
                        doc.add_paragraph(NO_DATA_DESCRIPTION)
                else:
                    doc.add_paragraph(str(value))
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)
