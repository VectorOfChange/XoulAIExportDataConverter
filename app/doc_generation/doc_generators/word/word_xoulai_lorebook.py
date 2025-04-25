# platform_specific_generators/word/word_xoulai_lorebook.py
from docx.document import Document as DocxDocument

from globals.globals import NO_DATA_DESCRIPTION
from models.platform_xoulai.lorebook_xoulai import LorebookSectionXoulAI, LorebookXoulAI

# Lorebook Helpers
def word_xoulai_add_lorebook_section_to_doc(doc: DocxDocument, section: LorebookSectionXoulAI):
    doc.add_heading(section.name or "Unnamed Section", level=5)

    # Iterate over the character fields dynamically
    for field_name, value in vars(section).items():
        if field_name == "name":
            continue  # already used above

        # Format the field name nicely (e.g., backstory_spec -> Backstory Spec)
        formatted_name = field_name.replace('_', ' ').title()
        doc.add_heading(formatted_name, level=6)

        # Handle the field content
        if value:
            if isinstance(value, list):
                doc.add_paragraph(", ".join(str(item) for item in value))
            else:
                doc.add_paragraph(str(value))
        else:
            doc.add_paragraph(NO_DATA_DESCRIPTION)

def word_xoulai_add_lorebook_to_doc(doc: DocxDocument, lorebook: LorebookXoulAI):
    doc.add_heading(lorebook.name or "Unnamed Lorebook", level=2)

    for field_name, value in vars(lorebook).items():
        if field_name == "name":
            continue  # Already used

        formatted_name = field_name.replace('_', ' ').title()

        if field_name == "embedded" and value:
            doc.add_heading("Embedded", level=3)
            doc.add_heading("Asset Type", level=4)
            doc.add_paragraph(value.asset_type or NO_DATA_DESCRIPTION)
            if value.sections:
                doc.add_heading("Sections", level=4)
                for section in value.sections:
                    word_xoulai_add_lorebook_section_to_doc(doc, section)
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)
        else:
            doc.add_heading(formatted_name, level=3)
            if value:
                if isinstance(value, list):
                    doc.add_paragraph(", ".join(str(item) for item in value))
                else:
                    doc.add_paragraph(str(value))
            else:
                doc.add_paragraph(NO_DATA_DESCRIPTION)


def word_xoulai_add_all_lorebooks_to_doc(doc: DocxDocument, lorebooks: list[LorebookXoulAI]):
    doc.add_page_break()
    doc.add_heading("Lorebooks", level=1)
    
    if lorebooks:
        for idx, lorebook in enumberate(lorebooks):
            word_xoulai_add_lorebook_to_doc(doc, lorebook)
            
            if idx < len(lorebooks) - 1:
                doc.add_page_break()
    else:
        doc.add_paragraph("No lorebooks were found in the Xoul Data Export ZIP file.")
