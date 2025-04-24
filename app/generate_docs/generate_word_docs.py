# generate_docs/generate_word_doc.py

from io import BytesIO
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement 
from models.all_data import AllData
from models.character import Character
from models.scenario import Scenario, PromptSpec, Objective

# Character Helpers
def add_character_to_doc(doc: DocxDocument, character: Character):
    # Add Heading 1 with the character's name
    doc.add_heading(character.name or "Unnamed Character", level=2)

    # Iterate over the character fields dynamically
    for field_name, value in vars(character).items():
        if field_name == "name":
            continue  # already used in Heading 1

        # Format the field name nicely (e.g., backstory_spec -> Backstory Spec)
        formatted_name = field_name.replace('_', ' ').title()
        doc.add_heading(formatted_name, level=3)

        # Handle the field content
        if value:
            if isinstance(value, list):
                doc.add_paragraph(", ".join(str(item) for item in value))
            else:
                doc.add_paragraph(str(value))
        else:
            doc.add_paragraph("None/Empty")


def add_all_characters_to_doc(doc: DocxDocument, characters):
    doc.add_heading("Xouls", level=1)
    if characters:
        for index, character in enumerate(characters):
            add_character_to_doc(doc, character)
            doc.add_page_break()
    else:
        doc.add_paragraph("No characters were found in the Xoul Data Export ZIP file.")

# Scenario Helpers
def add_scenario_to_doc(doc: DocxDocument, scenario: Scenario):
    doc.add_heading(scenario.name or "Unnamed Scenario", level=2)

    for field_name, value in vars(scenario).items():
        if field_name == "name":
            continue

        formatted_name = field_name.replace('_', ' ').title()
        if isinstance(value, PromptSpec):
            doc.add_heading("Prompt Spec", level=3)
            for sub_field, sub_value in vars(value).items():
                doc.add_heading(sub_field.replace('_', ' ').title(), level=4)
                doc.add_paragraph(str(sub_value) if sub_value else "None/Empty")

        elif isinstance(value, Objective):
            doc.add_heading("Objective", level=3)
            doc.add_heading("Description", level=4)
            doc.add_paragraph(value.description or "None/Empty")
            if value.meters:
                doc.add_heading("Meters", level=4)
                for meter in value.meters:
                    doc.add_heading(meter.name or "Unnamed Meter", level=5)
                    doc.add_paragraph(f"Description: {meter.description or 'None/Empty'}")
                    doc.add_paragraph(f"Value: {meter.value}")
            else:
                doc.add_paragraph("No meters")

        else:
            doc.add_heading(formatted_name, level=3)
            if value:
                if isinstance(value, list):
                    doc.add_paragraph(", ".join(str(item) for item in value))
                else:
                    doc.add_paragraph(str(value))
            else:
                doc.add_paragraph("None/Empty")

def add_all_scenarios_to_doc(doc: DocxDocument, scenarios):
    doc.add_heading("Scenarios", level=1)
    if scenarios:
        for index, scenario in enumerate(scenarios):
            add_scenario_to_doc(doc, scenario)
            if index < len(scenarios) - 1:
                doc.add_page_break()
    else:
        doc.add_paragraph("No scenarios were found in the Xoul Data Export ZIP file.")

def add_info_section_to_doc(doc: DocxDocument):
    doc.add_heading("Information About the Data", level=1)
    doc.add_heading("About this Document", level=2)
    doc.add_paragraph("This is the information that is within the Xoul Data Export ZIP file. Icon and picture links have been removed because they were just links to the Xoul AI site, which is now offline.")
    doc.add_heading("Slugs", level=2)
    doc.add_paragraph("A \"slug\" is a unique ID used to link data together. For example, if a scenario uses a lorebook, the scenario data will include the slug of the lorebook.")
    doc.add_paragraph("If you created the referenced item (for example, you want to go to the lorebook that was included in the scenario, and you created that lorebook), you can search this document for the slug and it should bring you to it.")
    doc.add_page_break()

# Generate Doc
def generate_word_docs(all_data: AllData) -> BytesIO:
    doc = Document()

    # Title Page
    title = doc.add_paragraph("Converted Xoul AI Data", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Xouls - Beta Test", style='Subtitle')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # TOC
    doc.add_heading("Table of Contents", 0)
    toc_paragraph = doc.add_paragraph()
    toc_run = toc_paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click on this text, then click Update Field to show the  Table of Contents."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = toc_run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    toc_run.add_break(WD_BREAK.PAGE)
    
    # Content
    add_info_section_to_doc(doc)
    add_all_characters_to_doc(doc, all_data.characters)
    add_all_scenarios_to_doc(doc, all_data.scenarios)
    

    # Save document to buffer
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    
    return doc_buffer

    # doc_filename = name.replace(".json", ".docx")
    # output_zip.writestr(doc_filename, doc_buffer.getvalue())
    # log(f"Word doc added to ZIP: {doc_filename}")